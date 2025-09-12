"""
Text Extraction Utilities for Chess Files.

This module handles text extraction from various file formats including
PDF, TXT, PGN, and other document types. It provides clean, readable text
that can be processed by the semantic analysis pipeline.
"""

import logging
from pathlib import Path
from typing import Optional

from ..core.constants import FileTypes

# Optional imports for different file types
try:
    import PyPDF2

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class TextExtractor:
    """
    Extract clean text from various file formats.

    Supports PDF, TXT, PGN, DOCX and other document formats commonly
    used for chess content. Handles encoding issues and provides
    clean text suitable for analysis.
    """

    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def extract_text(self, file_path: str) -> str:
        """
        Extract text content from file.

        Args:
            file_path: Path to the file to extract text from

        Returns:
            Extracted text content as string
        """
        try:
            file_path_obj = Path(file_path)
            file_ext = file_path_obj.suffix.lower()

            # Route to appropriate extraction method
            if file_ext in FileTypes.TEXT_EXTENSIONS:
                return self._extract_text_file(file_path)
            elif file_ext in FileTypes.PDF_EXTENSIONS:
                return self._extract_pdf(file_path)
            elif file_ext == '.docx':
                return self._extract_docx(file_path)
            elif file_ext == '.doc':
                return self._extract_doc_fallback(file_path)
            else:
                # Try to read as text file anyway
                return self._extract_text_file(file_path, fallback=True)

        except Exception as e:
            self.logger.error(f"Text extraction failed for {file_path}: {e}")
            return ""

    def _extract_text_file(self, file_path: str, fallback: bool = False) -> str:
        """
        Extract text from plain text files (TXT, PGN, etc.).

        Handles various encodings and provides clean text output.
        """
        encodings_to_try = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252', 'iso-8859-1']

        for encoding in encodings_to_try:
            try:
                with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()

                # Validate content quality
                if self._is_valid_text_content(content):
                    return self._clean_text_content(content)

            except (UnicodeDecodeError, UnicodeError) as e:
                if not fallback:
                    self.logger.debug(f"Encoding {encoding} failed for {file_path}: {e}")
                continue
            except Exception as e:
                if not fallback:
                    self.logger.warning(f"Failed to read {file_path} with {encoding}: {e}")
                continue

        if fallback:
            # Last resort: read first portion as binary and decode what we can
            try:
                with open(file_path, 'rb') as f:
                    raw_content = f.read(10000)  # First 10KB
                    content = raw_content.decode('utf-8', errors='ignore')
                    return self._clean_text_content(content)
            except:
                pass

        self.logger.warning(f"Could not extract text from {file_path}")
        return ""

    def _extract_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF files.

        Uses PyPDF2 for text extraction. Handles both text-based PDFs
        and provides meaningful error messages for scanned PDFs.
        """
        if not PDF_AVAILABLE:
            self.logger.warning(f"PyPDF2 not available for PDF extraction: {file_path}")
            return ""

        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text_parts = []

                total_pages = len(reader.pages)
                if total_pages == 0:
                    self.logger.warning(f"PDF has no pages: {file_path}")
                    return ""

                # Extract text from all pages
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    except Exception as e:
                        self.logger.debug(f"Failed to extract text from page {page_num} of {file_path}: {e}")
                        continue

                # Combine all text
                full_text = "\n".join(text_parts)

                # Validate extraction quality
                if len(full_text.strip()) < 100:
                    self.logger.info(f"PDF appears to be scanned or image-based: {file_path} "
                                     f"(extracted only {len(full_text.strip())} characters)")
                    return ""

                return self._clean_text_content(full_text)

        except Exception as e:
            self.logger.error(f"PDF extraction failed for {file_path}: {e}")
            return ""

    def _extract_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX files.
        """
        if not DOCX_AVAILABLE:
            self.logger.warning(f"python-docx not available for DOCX extraction: {file_path}")
            return ""

        try:
            doc = Document(file_path)
            text_parts = []

            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            full_text = "\n".join(text_parts)
            return self._clean_text_content(full_text)

        except Exception as e:
            self.logger.error(f"DOCX extraction failed for {file_path}: {e}")
            return ""

    def _extract_doc_fallback(self, file_path: str) -> str:
        """
        Fallback extraction for legacy DOC files.

        Since python-docx doesn't handle .doc files, we try to read
        what we can or suggest conversion.
        """
        self.logger.info(f"Legacy DOC file detected: {file_path}. "
                         f"Consider converting to DOCX for better extraction.")

        # Try to read as text (might get some readable content)
        return self._extract_text_file(file_path, fallback=True)

    def _is_valid_text_content(self, content: str) -> bool:
        """
        Validate that extracted content is meaningful text.

        Checks for minimum length, readable characters, and basic structure.
        """
        if not content or len(content.strip()) < 50:
            return False

        # Check for reasonable ratio of printable characters
        printable_chars = sum(1 for c in content if c.isprintable() or c.isspace())
        if printable_chars / len(content) < 0.7:
            return False

        # Check for some alphabetic content (not just numbers/symbols)
        alpha_chars = sum(1 for c in content if c.isalpha())
        if alpha_chars / len(content) < 0.1:
            return False

        return True

    def _clean_text_content(self, content: str) -> str:
        """
        Clean and normalize extracted text content.

        Handles common issues like extra whitespace, broken words,
        and encoding artifacts.
        """
        if not content:
            return ""

        # Handle unicode normalization
        content = content.replace('"', '"').replace('"', '"')
        content = content.replace(''', "'").replace(''', "'")
        content = content.replace('–', '-').replace('—', '-')

        # Fix common PDF extraction issues
        content = self._fix_pdf_artifacts(content)

        # Normalize whitespace
        content = ' '.join(content.split())

        # Remove excessive newlines but preserve paragraph structure
        content = '\n'.join(line.strip() for line in content.split('\n') if line.strip())

        return content.strip()

    def _fix_pdf_artifacts(self, text: str) -> str:
        """
        Fix common PDF text extraction artifacts.

        PDF extraction often breaks words across lines or creates
        other formatting issues that need cleanup.
        """
        import re

        # Fix broken words (letter space letter patterns)
        text = re.sub(r'(?<=[a-zA-Z])\s+(?=[a-zA-Z])\b(?![A-Z])', '', text)

        # Fix broken chess notation
        text = re.sub(r'(\d+)\s*\.\s*([KQRBN]?[a-h][1-8])', r'\1.\2', text)

        # Fix broken move sequences
        text = re.sub(r'([a-h][1-8])\s+([KQRBN]?[a-h][1-8])', r'\1 \2', text)

        # Remove excessive spaces
        text = re.sub(r'\s+', ' ', text)

        return text

    def extract_metadata(self, file_path: str) -> dict:
        """
        Extract metadata from files when available.

        Returns basic file information and any document-specific metadata.
        """
        metadata = {
            'file_size': 0,
            'file_type': '',
            'page_count': 0,
            'title': '',
            'author': '',
            'creation_date': '',
            'text_length': 0
        }

        try:
            file_path_obj = Path(file_path)
            file_stat = file_path_obj.stat()

            metadata['file_size'] = file_stat.st_size
            metadata['file_type'] = file_path_obj.suffix.lower()

            # Extract format-specific metadata
            if metadata['file_type'] == '.pdf' and PDF_AVAILABLE:
                metadata.update(self._extract_pdf_metadata(file_path))
            elif metadata['file_type'] == '.docx' and DOCX_AVAILABLE:
                metadata.update(self._extract_docx_metadata(file_path))

            # Get text length
            text_content = self.extract_text(file_path)
            metadata['text_length'] = len(text_content)

        except Exception as e:
            self.logger.debug(f"Metadata extraction failed for {file_path}: {e}")

        return metadata

    def _extract_pdf_metadata(self, file_path: str) -> dict:
        """Extract PDF-specific metadata"""
        metadata = {}

        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)

                metadata['page_count'] = len(reader.pages)

                # Extract document info if available
                if reader.metadata:
                    metadata['title'] = str(reader.metadata.get('/Title', '')).strip()
                    metadata['author'] = str(reader.metadata.get('/Author', '')).strip()
                    metadata['creation_date'] = str(reader.metadata.get('/CreationDate', '')).strip()

        except Exception as e:
            self.logger.debug(f"PDF metadata extraction failed: {e}")

        return metadata

    def _extract_docx_metadata(self, file_path: str) -> dict:
        """Extract DOCX-specific metadata"""
        metadata = {}

        try:
            doc = Document(file_path)

            # Count paragraphs as pages approximation
            metadata['page_count'] = max(1, len(doc.paragraphs) // 20)

            # Extract core properties
            core_props = doc.core_properties
            metadata['title'] = core_props.title or ''
            metadata['author'] = core_props.author or ''
            if core_props.created:
                metadata['creation_date'] = core_props.created.isoformat()

        except Exception as e:
            self.logger.debug(f"DOCX metadata extraction failed: {e}")

        return metadata

    def is_supported_format(self, file_path: str) -> bool:
        """
        Check if file format is supported for text extraction.

        Args:
            file_path: Path to check

        Returns:
            True if format is supported
        """
        file_ext = Path(file_path).suffix.lower()
        supported_extensions = FileTypes.CHESS_EXTENSIONS

        # Check library availability for specific formats
        if file_ext == '.pdf' and not PDF_AVAILABLE:
            return False
        elif file_ext == '.docx' and not DOCX_AVAILABLE:
            return False

        return file_ext in supported_extensions

    def estimate_extraction_time(self, file_path: str) -> float:
        """
        Estimate time required for text extraction.

        Useful for progress tracking and batch processing planning.

        Returns:
            Estimated extraction time in seconds
        """
        try:
            file_size = Path(file_path).stat().st_size
            file_ext = Path(file_path).suffix.lower()

            # Time estimates based on file type and size (in MB)
            size_mb = file_size / (1024 * 1024)

            if file_ext in ['.txt', '.pgn']:
                return max(0.1, size_mb * 0.05)  # Very fast for text files
            elif file_ext == '.pdf':
                return max(0.5, size_mb * 0.3)  # Slower for PDFs
            elif file_ext == '.docx':
                return max(0.3, size_mb * 0.2)  # Moderate for DOCX
            else:
                return max(0.2, size_mb * 0.1)  # Default estimate

        except Exception:
            return 1.0  # Default fallback